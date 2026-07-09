#!/usr/bin/env python3
# =============================================================================
#  verificar-temario.py — coherencia FUENTE (.md) <-> BUILD (.docx) del temario
# -----------------------------------------------------------------------------
#  Implementa las cinco comprobaciones (a-e) de la SPEC-003 §3.3, rescatadas a
#  tooling versionado por la SPEC-004 §4.1.
#
#  ¿Por qué Python, si el andamiaje del curso es 100 % Bash?
#  La regla "prohibido Python" aplica al `bin/` que ejecuta el ALUMNO, por
#  portabilidad macOS/Git Bash. Esto es `tools/`: tooling INTERNO del repo, que
#  corre en CI sobre ubuntu-latest. Python es válido aquí. La distinción es la
#  misma que hace el eslint.config.mjs del curso de Cypress (P-01 del ADN):
#  el tooling del repo no es material del curso.
#
#  Lecciones ya incorporadas (discrepancias reportadas en la SPEC-003):
#   · D1 — una fila válida se reconoce por su PRIMERA CELDA NUMÉRICA, no por la
#     forma del texto (negritas). Confundir forma con propiedad es el defecto
#     A-01 del ADN; la primera versión de este script cayó en él.
#   · D2 — las tablas del .docx se leen como CELDAS con python-docx, no como
#     texto renderizado: `pandoc -t plain` parte los nombres largos en varias
#     líneas y obliga a adivinar con regex.
#
#  Salida: greppeable, sin colores ANSI.  exit 0 = todo cuadra · exit 1 = no.
#  Uso:  python3 tools/verificar-temario.py
# =============================================================================
import re
import sys
from decimal import Decimal
from pathlib import Path

import docx

RAIZ = Path(__file__).resolve().parent.parent
MD = RAIZ / "docs/temario/TEMARIO-SPRING-BOOT-SII-v3.md"
DOCX = RAIZ / "docs/temario/TEMARIO-SPRING-BOOT-SII-v3.docx"

TOTAL = Decimal("36.0")
N_MODULOS = 15
N_SESIONES = 12
N_CAMBIOS = 13

RE_MOD_MD = re.compile(r"^\|\s*\*\*(\d+)\*\*\s*\|[^|]+\|\s*(\d,\d)\s*\|", re.M)
RE_SES_MD = re.compile(r"^\|\s*(S\d\d)\s*\|\s*\d+\s*\|\s*(.+?)\s*\|\s*(\d,\d)\s*\|", re.M)
RE_ASIG = re.compile(r"M(\d+)\s*\((\d,\d)\)")
RE_FILA_NUM_MD = re.compile(r"^\|\s*(\d+)\s*\|")

fallos = []


def num(s):
    return Decimal(s.strip().replace(",", "."))


def fmt(d):
    return f"{d:.1f}".replace(".", ",")


def check(cond, etiqueta):
    print(f"  [{'OK' if cond else 'ERROR'}] {etiqueta}")
    if not cond:
        fallos.append(etiqueta)
    return cond


def texto_docx(doc):
    """Todo el texto del .docx: párrafos + celdas. Sin pandoc: CI solo instala python-docx."""
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for fila in t.rows:
            partes.extend(c.text for c in fila.cells)
    return "\n".join(partes)


def main():
    for ruta in (MD, DOCX):
        if not ruta.exists():
            print(f"  [ERROR] No existe: {ruta.relative_to(RAIZ)}")
            return 1

    md = MD.read_text(encoding="utf-8")
    doc = docx.Document(DOCX)
    plano = texto_docx(doc)

    # --- extracción .md ---
    md_mod = {int(m[0]): num(m[1]) for m in RE_MOD_MD.findall(md)}
    md_ses = {}
    for sesion, celda, total in RE_SES_MD.findall(md):
        md_ses[sesion] = ([(int(a), num(b)) for a, b in RE_ASIG.findall(celda)], num(total))

    # --- extracción .docx (celdas reales, no texto renderizado: D2) ---
    dx_mod, dx_ses = {}, {}
    for t in doc.tables:
        for fila in t.rows:
            c = [x.text.strip() for x in fila.cells]
            if len(c) >= 3 and re.fullmatch(r"\d+", c[0]) and re.fullmatch(r"\d,\d", c[2]):
                dx_mod[int(c[0])] = num(c[2])
            if len(c) >= 4 and re.fullmatch(r"S\d\d", c[0]):
                dx_ses[c[0]] = ([(int(a), num(b)) for a, b in RE_ASIG.findall(c[2])], num(c[3]))

    esperado = [f"S{i:02d}" for i in range(1, N_SESIONES + 1)]

    # --- (a) versiones ---
    print("\n(a) Versiones declaradas en AMBOS")
    for token in ("Java 25 LTS", "Spring Boot 4.1"):
        check(token in md and token in plano, f'"{token}" presente en .md y en .docx')

    # --- (b) 15 módulos y total 36,0 ---
    print(f"\n(b) {N_MODULOS} modulos y total {fmt(TOTAL)}")
    rango = list(range(1, N_MODULOS + 1))
    check(sorted(md_mod) == rango, f".md declara los modulos 1..{N_MODULOS} ({len(md_mod)} hallados)")
    check(sorted(dx_mod) == rango, f".docx declara los modulos 1..{N_MODULOS} ({len(dx_mod)} hallados)")
    check(md_mod == dx_mod, "las horas por modulo COINCIDEN entre .md y .docx")
    check(sum(md_mod.values()) == TOTAL, f"suma de los modulos en .md = {fmt(sum(md_mod.values()))}")
    check(sum(dx_mod.values()) == TOTAL, f"suma de los modulos en .docx = {fmt(sum(dx_mod.values()))}")
    check(fmt(TOTAL) in md and fmt(TOTAL) in plano, f'literal "{fmt(TOTAL)}" presente en ambos')

    # --- (c) matriz de 12 filas ---
    print(f"\n(c) Matriz Modulo x Sesion con {N_SESIONES} filas S01..S{N_SESIONES}")
    check(sorted(md_ses) == esperado, f".md: {len(md_ses)} filas")
    check(sorted(dx_ses) == esperado, f".docx: {len(dx_ses)} filas")
    check(md_ses == dx_ses, "la matriz COINCIDE celda a celda entre .md y .docx")

    if sorted(md_ses) != esperado:
        print("\n  [ERROR] Sin matriz completa no se puede verificar la cuadratura (d).")
        return 1

    # --- (d) cuadratura: cálculo real, por las dos vías ---
    print("\n(d) CUADRATURA - calculo real\n")
    print("  Por sesion: suma de asignaciones vs Total declarado")
    total_ses = Decimal(0)
    for s in esperado:
        asigs, declarado = md_ses[s]
        suma = sum(h for _, h in asigs)
        total_ses += declarado
        detalle = " + ".join(f"M{m} ({fmt(h)})" for m, h in asigs)
        ok = suma == declarado
        print(f"    {s}  {detalle:<40} = {fmt(suma)}  decl {fmt(declarado)}  {'OK' if ok else 'ERROR'}")
        if not ok:
            fallos.append(f"la sesion {s} no cuadra")
    print(f"\n    TOTAL {N_SESIONES} sesiones = {fmt(total_ses)}")
    check(total_ses == TOTAL, f"suma por sesiones = {fmt(TOTAL)}")

    print("\n  Por modulo: horas repartidas en la matriz vs columna Horas de la tabla")
    reparto = {}
    for s in esperado:
        for m, h in md_ses[s][0]:
            reparto.setdefault(m, []).append((s, h))
    for m in rango:
        trozos = reparto.get(m, [])
        suma = sum(h for _, h in trozos)
        declarado = md_mod[m]
        detalle = " + ".join(f"{fmt(h)} [{s}]" for s, h in trozos)
        ok = suma == declarado
        print(f"    M{m:<2} {detalle:<34} = {fmt(suma)}  tabla {fmt(declarado)}  {'OK' if ok else 'ERROR'}")
        if not ok:
            fallos.append(f"el modulo M{m} no cuadra")
    repartido = sum(sum(h for _, h in v) for v in reparto.values())
    print(f"\n    TOTAL repartido en la matriz = {fmt(repartido)}")
    check(repartido == TOTAL, f"suma de horas repartidas = {fmt(TOTAL)}")
    check(set(reparto) == set(rango), "los 15 modulos aparecen en la matriz")

    # --- (e) anexo con 13 cambios ---
    #  Fila valida = primera celda numerica, en ambos formatos (leccion D1).
    print(f"\n(e) Anexo: {N_CAMBIOS} filas de cambios")
    anexo_md = md.split("## Anexo")[-1]
    nums_md = [int(mm.group(1)) for l in anexo_md.splitlines()
               if (mm := RE_FILA_NUM_MD.match(l.strip()))]
    nums_dx = [int(f.cells[0].text.strip()) for f in doc.tables[-1].rows
               if re.fullmatch(r"\d+", f.cells[0].text.strip())]
    objetivo = list(range(1, N_CAMBIOS + 1))
    check(nums_md == objetivo, f".md: cambios numerados 1..{N_CAMBIOS} sin huecos ({len(nums_md)} filas)")
    check(nums_dx == objetivo, f".docx: cambios numerados 1..{N_CAMBIOS} sin huecos ({len(nums_dx)} filas)")

    # --- veredicto ---
    print("\n" + "=" * 64)
    if fallos:
        print(f"VEREDICTO: {len(fallos)} ERROR(ES) - el temario NO es coherente")
        for f in fallos:
            print(f"  - {f}")
        print("\nRecuerda: ante divergencia, manda el .md (docs/temario/README.md).")
        return 1
    print("VEREDICTO: las 5 verificaciones PASAN. El .md y el .docx son coherentes.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
